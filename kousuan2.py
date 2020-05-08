#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Apr 20 14:42:36 2020

@author: qiang.wei
"""

import random
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import time


symbol = ['add', 'subtract', 'multiply', 'divide', 'divide2']
def add(maxsum = 100):
    first = random.randint(30, maxsum - 10)
    sum_ = random.randint(first + 10, maxsum)
    second = sum_ - first
    line = str(first) + '+' + str(second) + '='
    return line

def subtract(maxsub = 100):
    second = random.randint(30, maxsub-5)
    first = random.randint(second+5, maxsub)
    line = str(first) + '-' + str(second) + '='
    return line


def multiply():
    second = random.randint(2, 9)
    first = random.randint(2, 9)
    line = str(first) + 'x' + str(second) + '='
    return line

def divide():
    second = random.randint(5, 9)
    first = random.randint(2, 9)*second
    line = str(first) + '÷' + str(second) + '='
    return line

def divide2():
    second = random.randint(5, 9)
    first = random.randint(1, 9)*second + random.randint(1, second-1)
    line = str(first) + '÷' + str(second) + '='
    return line

def main():   
    switch = {'add': add,
              'subtract': subtract,
              'multiply': multiply,
              'divide': divide,
              'divide2': divide2
              }
    lines = []
    i = 0
    while i < 20:
        item=[]
        j = 0
        while j < 5:
            a = random.randint(0,4)
            item.append(switch.get(symbol[a], divide2)())
            j = j + 1
        lines.append(item)
        i = i + 1
    matrix = np.array(lines)

    doc = Document()
    paragraph = doc.add_heading('',level=0)  
    run = paragraph.add_run(text="小学生口算练习题", style=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size= Pt(20)
    time_ = time.asctime( time.localtime(time.time()) )
    print(time_)
    paragraph.add_run(text=time_, style=None).font.size= Pt(10)
    table = doc.add_table(matrix.shape[0], matrix.shape[1], style=None)
    for i in range(matrix.shape[0]):
        j = 0
        for j in range(matrix.shape[1]):
            table.cell(i, j).paragraphs[0].add_run(text=str(lines[i][j]), style=None).font.size= Pt(14)
    doc.save('小学生口算练习题.docx')

if __name__ == '__main__':
    main()
    